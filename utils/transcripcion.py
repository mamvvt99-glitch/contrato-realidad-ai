import streamlit as st
import os
import tempfile
import whisper
from datetime import datetime
import io
from docx import Document

def cargar_modelo_whisper():
    """
    Carga el modelo Whisper para transcripción.
    """
    try:
        import ssl
        import certifi
        import os
        
        # Configurar certificados SSL para macOS
        ssl_context = ssl.create_default_context(cafile=certifi.where())
        
        # Configurar variables de entorno para SSL
        os.environ['REQUESTS_CA_BUNDLE'] = certifi.where()
        os.environ['SSL_CERT_FILE'] = certifi.where()
        
        # Usar modelo base para velocidad, cambiar a 'medium' o 'large' para mejor precisión
        model = whisper.load_model("base")
        return model
    except Exception as e:
        st.error(f"Error al cargar el modelo Whisper: {str(e)}")
        st.info("💡 Solución: Ejecuta 'pip install certifi' y reinicia la aplicación")
        return None

def cargar_modelo_whisper_alternativo():
    """
    Carga el modelo Whisper usando un enfoque alternativo para evitar problemas de certificados.
    """
    try:
        import os
        import ssl
        
        # Deshabilitar verificación SSL temporalmente (solo para descarga del modelo)
        ssl._create_default_https_context = ssl._create_unverified_context
        
        # Usar modelo base para velocidad
        model = whisper.load_model("base")
        return model
    except Exception as e:
        st.error(f"Error al cargar el modelo Whisper (método alternativo): {str(e)}")
        return None

def transcribir_audio(audio_file, modelo):
    """
    Transcribe un archivo de audio usando Whisper.
    """
    try:
        # Crear archivo temporal
        with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp_file:
            tmp_file.write(audio_file.read())
            tmp_path = tmp_file.name
        
        # Transcribir
        result = modelo.transcribe(tmp_path)
        
        # Limpiar archivo temporal
        os.unlink(tmp_path)
        
        return result["text"]
    except Exception as e:
        st.error(f"Error en la transcripción: {str(e)}")
        return None

def render_transcripcion_inline():
    """
    Renderiza la interfaz de transcripción para usar en el flujo principal.
    
    Returns:
        Texto transcrito o None
    """
    st.markdown("### 🎤 Transcripción de Entrevista")
    
    st.info("""
    **Formatos soportados:** MP3, WAV, M4A, FLAC, OGG  
    **Tamaño máximo:** 25 MB
    """)
    
    # Cargar modelo Whisper
    with st.spinner("Cargando modelo de transcripción..."):
        modelo = cargar_modelo_whisper()
        
        # Si falla, intentar método alternativo
        if modelo is None:
            st.warning("⚠️ Intentando método alternativo...")
            modelo = cargar_modelo_whisper_alternativo()
    
    if modelo is None:
        st.error("❌ No se pudo cargar el modelo de transcripción.")
        st.info("💡 Soluciones:")
        st.info("1. Verifica tu conexión a internet")
        st.info("2. Ejecuta: pip install --upgrade certifi")
        st.info("3. Reinicia la aplicación")
        return None
    
    st.success("✅ Modelo de transcripción cargado correctamente")
    
    # Subir archivo de audio
    uploaded_file = st.file_uploader(
        "Selecciona un archivo de audio:",
        type=['mp3', 'wav', 'm4a', 'flac', 'ogg'],
        help="Formatos soportados: MP3, WAV, M4A, FLAC, OGG"
    )
    
    texto_transcrito = None
    
    if uploaded_file is not None:
        # Mostrar información del archivo
        file_details = {
            "Nombre": uploaded_file.name,
            "Tipo": uploaded_file.type,
            "Tamaño": f"{uploaded_file.size / 1024 / 1024:.2f} MB"
        }
        
        col1, col2, col3 = st.columns(3)
        for idx, (key, value) in enumerate(file_details.items()):
            with [col1, col2, col3][idx]:
                st.metric(key, value)
        
        # Botón para transcribir
        if st.button("🎤 Transcribir Audio", type="primary", use_container_width=True):
            if uploaded_file.size > 25 * 1024 * 1024:  # 25 MB
                st.error("❌ El archivo es demasiado grande. El tamaño máximo es 25 MB.")
            else:
                with st.spinner("Transcribiendo audio... Esto puede tomar unos minutos."):
                    # Transcribir
                    texto_transcrito = transcribir_audio(uploaded_file, modelo)
                    
                    if texto_transcrito:
                        st.success("✅ Transcripción completada exitosamente!")
                        
                        # Mostrar transcripción
                        st.text_area(
                            "📝 Transcripción:",
                            value=texto_transcrito,
                            height=300,
                            help="Puedes editar el texto si es necesario"
                        )
                        
                        # Estadísticas
                        palabras = len(texto_transcrito.split())
                        caracteres = len(texto_transcrito)
                        
                        col1, col2, col3 = st.columns(3)
                        with col1:
                            st.metric("Palabras", palabras)
                        with col2:
                            st.metric("Caracteres", caracteres)
                        with col3:
                            st.metric("Tiempo estimado", f"{palabras/150:.1f} min")
    
    return texto_transcrito

def generar_documento_transcripcion(texto_transcripcion, nombre_archivo, fecha_transcripcion):
    """
    Genera un documento Word con la transcripción.
    """
    doc = Document()
    
    # Título
    doc.add_heading("Transcripción de Entrevista", level=1)
    
    # Información del archivo
    doc.add_heading("Información del Archivo", level=2)
    doc.add_paragraph(f"Archivo original: {nombre_archivo}")
    doc.add_paragraph(f"Fecha de transcripción: {fecha_transcripcion}")
    doc.add_paragraph(f"Modelo utilizado: Whisper Base")
    
    # Separador
    doc.add_paragraph("")
    
    # Transcripción
    doc.add_heading("Transcripción", level=2)
    doc.add_paragraph(texto_transcripcion)
    
    return doc

def render_transcripcion_module():
    """
    Renderiza el módulo de transcripción en Streamlit.
    """
    st.markdown("""
    <div style="text-align: center; padding: 2rem 0; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); 
                border-radius: 15px; margin-bottom: 2rem;">
        <h1 style="color: white; margin: 0;">🎤 Transcripción de Audios</h1>
        <p style="color: white; font-size: 1.1rem; margin-top: 0.5rem;">Convierte audio en texto con Inteligencia Artificial</p>
    </div>
    """, unsafe_allow_html=True)
    
    st.info("""
    **¿Qué hace este módulo?**
    
    Permite transcribir archivos de audio de entrevistas con clientes, testigos o cualquier 
    conversación legal relevante. Utiliza inteligencia artificial para convertir el habla en texto 
    de manera precisa y rápida.
    
    **Formatos soportados:** MP3, WAV, M4A, FLAC, OGG  
    **Tamaño máximo:** 25 MB
    """)
    
    st.markdown("---")
    
    # Cargar modelo Whisper
    with st.spinner("Cargando modelo de transcripción..."):
        modelo = cargar_modelo_whisper()
        
        # Si falla, intentar método alternativo
        if modelo is None:
            st.warning("⚠️ Intentando método alternativo...")
            modelo = cargar_modelo_whisper_alternativo()
    
    if modelo is None:
        st.error("❌ No se pudo cargar el modelo de transcripción.")
        st.info("💡 Soluciones:")
        st.info("1. Verifica tu conexión a internet")
        st.info("2. Ejecuta: pip install --upgrade certifi")
        st.info("3. Reinicia la aplicación")
        return
    
    st.success("✅ Modelo de transcripción cargado correctamente")
    
    # Subir archivo de audio
    st.markdown("### 📁 Subir Archivo de Audio")
    
    uploaded_file = st.file_uploader(
        "Selecciona un archivo de audio:",
        type=['mp3', 'wav', 'm4a', 'flac', 'ogg'],
        help="Formatos soportados: MP3, WAV, M4A, FLAC, OGG"
    )
    
    if uploaded_file is not None:
        # Mostrar información del archivo
        file_details = {
            "Nombre del archivo": uploaded_file.name,
            "Tipo de archivo": uploaded_file.type,
            "Tamaño": f"{uploaded_file.size / 1024 / 1024:.2f} MB"
        }
        
        st.subheader("📋 Información del Archivo")
        for key, value in file_details.items():
            st.write(f"**{key}:** {value}")
        
        # Botón para transcribir
        if st.button("🎤 Transcribir Audio", type="primary", use_container_width=True):
            if uploaded_file.size > 25 * 1024 * 1024:  # 25 MB
                st.error("❌ El archivo es demasiado grande. El tamaño máximo es 25 MB.")
            else:
                with st.spinner("Transcribiendo audio... Esto puede tomar unos minutos."):
                    # Transcribir
                    transcripcion = transcribir_audio(uploaded_file, modelo)
                    
                    if transcripcion:
                        st.success("✅ Transcripción completada exitosamente!")
                        
                        # Mostrar transcripción
                        st.subheader("📝 Transcripción")
                        st.text_area(
                            "Texto transcrito:",
                            value=transcripcion,
                            height=300,
                            help="Puedes editar el texto si es necesario"
                        )
                        
                        # Estadísticas
                        palabras = len(transcripcion.split())
                        caracteres = len(transcripcion)
                        
                        col1, col2, col3 = st.columns(3)
                        with col1:
                            st.metric("Palabras", palabras)
                        with col2:
                            st.metric("Caracteres", caracteres)
                        with col3:
                            st.metric("Tiempo estimado", f"{palabras/150:.1f} min")
                        
                        # Botones de exportación
                        st.subheader("📥 Exportar Transcripción")
                        
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            # Exportar como TXT
                            st.download_button(
                                label="📄 Descargar como TXT",
                                data=transcripcion,
                                file_name=f"transcripcion_{uploaded_file.name.split('.')[0]}.txt",
                                mime="text/plain"
                            )
                        
                        with col2:
                            # Exportar como DOCX
                            doc = generar_documento_transcripcion(
                                transcripcion,
                                uploaded_file.name,
                                datetime.now().strftime("%d/%m/%Y %H:%M")
                            )
                            
                            buffer = io.BytesIO()
                            doc.save(buffer)
                            buffer.seek(0)
                            
                            st.download_button(
                                label="📄 Descargar como DOCX",
                                data=buffer,
                                file_name=f"transcripcion_{uploaded_file.name.split('.')[0]}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        
                        # Funcionalidades adicionales
                        st.subheader("🔧 Funcionalidades Adicionales")
                        
                        # Resumen automático
                        if st.button("📋 Generar Resumen Automático"):
                            with st.spinner("Generando resumen..."):
                                # Aquí podrías integrar con OpenAI para generar un resumen
                                st.info("Funcionalidad de resumen automático próximamente disponible.")
                        
                        # Análisis de sentimientos
                        if st.button("😊 Análisis de Sentimientos"):
                            with st.spinner("Analizando sentimientos..."):
                                # Aquí podrías integrar análisis de sentimientos
                                st.info("Funcionalidad de análisis de sentimientos próximamente disponible.")
                        
                        # Extracción de puntos clave
                        if st.button("🎯 Extraer Puntos Clave"):
                            with st.spinner("Extrayendo puntos clave..."):
                                # Aquí podrías integrar extracción de puntos clave
                                st.info("Funcionalidad de extracción de puntos clave próximamente disponible.")
                        
                    else:
                        st.error("❌ Error en la transcripción. Verifica que el archivo sea válido.")
    
    # Información adicional
    st.markdown("---")
    st.subheader("ℹ️ Información Técnica")
    
    with st.expander("Ver detalles técnicos"):
        st.markdown("""
        **Modelo utilizado:** Whisper Base (OpenAI)
        
        **Características:**
        - Reconocimiento de voz en español
        - Procesamiento local (no se envían datos a servidores externos)
        - Alta precisión en transcripciones
        - Soporte para múltiples acentos
        
        **Limitaciones:**
        - Tamaño máximo de archivo: 25 MB
        - Tiempo de procesamiento depende del tamaño del archivo
        - Mejor rendimiento con audio de buena calidad
        
        **Recomendaciones:**
        - Usar archivos de audio con buena calidad
        - Evitar archivos con mucho ruido de fondo
        - Para archivos largos, considera dividirlos en segmentos más pequeños
        """)
    
    # Consejos de uso
    st.subheader("💡 Consejos para Mejor Transcripción")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("""
        **✅ Hacer:**
        - Usar audio de buena calidad
        - Hablar claramente
        - Minimizar ruido de fondo
        - Usar archivos en formato WAV o MP3
        """)
    
    with col2:
        st.markdown("""
        **❌ Evitar:**
        - Audio con mucho eco
        - Múltiples personas hablando simultáneamente
        - Archivos muy comprimidos
        - Audio con música de fondo
        """) 