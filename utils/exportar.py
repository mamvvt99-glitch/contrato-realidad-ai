# utils/exportar.py

from docx import Document
from datetime import datetime

def generar_docx_concepto(hechos: str, concepto: str) -> Document:
    doc = Document()

    doc.add_heading('Concepto Jurídico de Viabilidad', level=1)
    doc.add_paragraph(f"Fecha de emisión: {datetime.now().strftime('%d/%m/%Y %H:%M')}")


    doc.add_heading('Concepto Jurídico', level=2)
    doc.add_paragraph(concepto)

    return doc