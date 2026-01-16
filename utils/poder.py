import streamlit as st
from docx import Document
from docx.shared import Inches
import io
import re

def extraer_campos_poder():
    """
    Extrae los campos reemplazables del documento de poder de referencia.
    Retorna un diccionario con los campos y sus valores por defecto.
    """
    campos = {
        "juez_administrativo": "JUEZ ADMINISTRATIVO DEL CIRCUITO (reparto) E.S.D.",
        "nombre_poderdante": "JULIE MELISSA MORENO ALARCÓN",
        "cedula_poderdante": "52.982.304",
        "ciudad_poderdante": "Bogotá",
        "nombre_organizacion": "CONDE ABOGADOS ASOCIADOS SAS",
        "nit_organizacion": "828002664-3",
        "fecha_constitucion": "26 de marzo de 2014",
        "numero_libro": "00007635",
        "libro": "lX",
        "certificado_existencia": "2021",
        "nombre_representante": "MARCELA PATRICIA CEBALLOS OSORIO",
        "cedula_representante": "1.075.227.003",
        "ciudad_representante": "Neiva",
        "tarjeta_profesional": "214.303",
        "consejo_judicatura": "Consejo Superior de la Judicatura",
        "entidad_demandada": "LA NACIÓN – MINISTERIO DE DEFENSA -POLICIA NACIONAL – DIRECCIÓN DE SANIDAD",
        "oficio_1": "S-2020-465025-HEBOG/RASES-GRUCO 29.25",
        "fecha_oficio_1": "29 de diciembre de 2020",
        "jefe_regional": "Jefe Regional de Aseguramiento en Salud No. 1",
        "oficio_2": "S-2021-000811-DISAN ASJUR -41.10",
        "fecha_oficio_2": "07 de enero del 2021",
        "directora_sanidad": "Directora de Sanidad de la Policía Nacional",
        "cargo_laboral": "Jefe de Enfermería",
        "fecha_inicio_laboral": "enero del 2011",
        "fecha_fin_laboral": "26 de septiembre de 2019",
        "ciudad_poderdante_firma": "Bogotá D.C",
        "ciudad_representante_firma": "Neiva"
    }
    return campos

def generar_documento_poder(campos_llenados):
    """
    Genera un nuevo documento de poder con los campos llenados.
    """
    # Plantilla del documento de poder
    plantilla = f"""
Señores {campos_llenados['juez_administrativo']}

{campos_llenados['nombre_poderdante']}, identificada con la cédula de ciudadanía número {campos_llenados['cedula_poderdante']} de {campos_llenados['ciudad_poderdante']}, de forma atenta y respetuosa me permito manifestar a usted, que confiero poder especial, amplio y suficiente a la Organización Jurídica {campos_llenados['nombre_organizacion']}, con NIT No {campos_llenados['nit_organizacion']}, constituida por documento privado en junta de socios del {campos_llenados['fecha_constitucion']}, bajo el número {campos_llenados['numero_libro']} del libro {campos_llenados['libro']}, según consta en el certificado de existencia y representación legal del {campos_llenados['certificado_existencia']}, y cuyo representante legal es la Doctora {campos_llenados['nombre_representante']}, identificada con cedula de ciudadanía No. {campos_llenados['cedula_representante']} de {campos_llenados['ciudad_representante']} y tarjeta profesional No. {campos_llenados['tarjeta_profesional']} del {campos_llenados['consejo_judicatura']}, para que, en mi nombre y representación, presente medio de control de Nulidad y Restablecimiento del Derecho contra {campos_llenados['entidad_demandada']}, a fin de que de declare la nulidad de los efectos económicos de los oficios No. {campos_llenados['oficio_1']} de {campos_llenados['fecha_oficio_1']} expedido por el {campos_llenados['jefe_regional']} y el oficio {campos_llenados['oficio_2']} del {campos_llenados['fecha_oficio_2']}, expedido por la {campos_llenados['directora_sanidad']}, por medio del cual se niega el reconocimiento del vínculo laboral y pago de emolumentos salariales, prestacionales, indemnizatorios y de seguridad social a la suscrita y a título de restablecimiento del derecho se ordene la existencia del vínculo laboral entre dicha entidad y la suscrita, la cual tuvo vigencia entre {campos_llenados['fecha_inicio_laboral']} hasta el {campos_llenados['fecha_fin_laboral']}, tiempo durante el cual me desempeñé como {campos_llenados['cargo_laboral']} para esta entidad, como consecuencia de lo anterior se ordene cancelar la diferencia existente entre la suscrita como {campos_llenados['cargo_laboral']} y, un {campos_llenados['cargo_laboral']} de planta y/o uno uniformado de conformidad con lo que se pagaba desde el año 2011 hasta el año 2019 a dicho cargo de planta, así mismo se ordene el reconocimiento y pago de todos los emolumentos prestacionales, indemnizatorios y de seguridad social (en la proporción patronal) que se dejaron de cancelar durante toda la relación laboral, de conformidad con lo devengado por un {campos_llenados['cargo_laboral']} de planta y con lo que se pagaba desde el año 2011 y hasta el año 2019 a dicho cargo de planta, que se ordene la devolución de toda deducción existente durante la relación laboral, y se ordene y pague las indemnizaciones por la no consignación de cesantías al fondo de cesantías, sobre las sumas adeudadas se ordene la correspondiente indexación y el pago de los intereses moratorios y se condene en costas a la demandada.

El apoderado especial queda facultado para recibir, transigir, desistir, sustituir, reasumir, conciliar, renunciar; además, facultad expresa de cobrar y recibir el pago de los reconocimientos que así se pretenden, y en general todas las demás facultades necesarias para el cumplimiento de este mandato y conforme a lo establecido por el articulo 77 del Código General del Proceso.

Respetuosamente,

{campos_llenados['nombre_poderdante']} C.C. {campos_llenados['cedula_poderdante']} de {campos_llenados['ciudad_poderdante_firma']}

Acepto,

{campos_llenados['nombre_representante']}
Representante Legal {campos_llenados['nombre_organizacion']} C.C. {campos_llenados['cedula_representante']} de {campos_llenados['ciudad_representante_firma']} T.P. No {campos_llenados['tarjeta_profesional']} del C. S. de la J.
"""
    
    # Crear documento Word
    doc = Document()
    
    # Agregar contenido al documento
    parrafos = plantilla.strip().split('\n\n')
    for parrafo in parrafos:
        if parrafo.strip():
            doc.add_paragraph(parrafo.strip())
    
    return doc

def render_poder_module():
    """
    Renderiza el módulo de poder en Streamlit.
    """
    st.markdown("""
    <div style="text-align: center; padding: 2rem 0; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); 
                border-radius: 15px; margin-bottom: 2rem;">
        <h1 style="color: white; margin: 0;">📋 Módulo de Poder Especial</h1>
        <p style="color: white; font-size: 1.1rem; margin-top: 0.5rem;">Generación de Poder Especial para Representación Legal</p>
    </div>
    """, unsafe_allow_html=True)
    
    st.info("""
    **¿Qué es un Poder Especial?**
    
    Un poder especial es un documento legal que autoriza a un abogado o firma jurídica 
    para representar a un cliente en un proceso judicial específico. Complete los siguientes 
    campos para generar el documento automáticamente.
    """)
    
    st.markdown("---")
    
    # Extraer campos del documento de referencia
    campos = extraer_campos_poder()
    
    st.markdown("### 📝 Información del Poderdante")
    st.markdown("---")
    col1, col2 = st.columns(2)
    
    with col1:
        campos['nombre_poderdante'] = st.text_input(
            "Nombre completo del poderdante:",
            value=campos['nombre_poderdante']
        )
        campos['cedula_poderdante'] = st.text_input(
            "Cédula del poderdante:",
            value=campos['cedula_poderdante']
        )
        campos['ciudad_poderdante'] = st.text_input(
            "Ciudad del poderdante:",
            value=campos['ciudad_poderdante']
        )
    
    with col2:
        campos['juez_administrativo'] = st.text_input(
            "Juez Administrativo:",
            value=campos['juez_administrativo']
        )
        campos['ciudad_poderdante_firma'] = st.text_input(
            "Ciudad para firma:",
            value=campos['ciudad_poderdante_firma']
        )
    
    st.markdown("### 🏢 Información de la Organización")
    st.markdown("---")
    col1, col2 = st.columns(2)
    
    with col1:
        campos['nombre_organizacion'] = st.text_input(
            "Nombre de la organización:",
            value=campos['nombre_organizacion']
        )
        campos['nit_organizacion'] = st.text_input(
            "NIT de la organización:",
            value=campos['nit_organizacion']
        )
        campos['fecha_constitucion'] = st.text_input(
            "Fecha de constitución:",
            value=campos['fecha_constitucion']
        )
    
    with col2:
        campos['numero_libro'] = st.text_input(
            "Número del libro:",
            value=campos['numero_libro']
        )
        campos['libro'] = st.text_input(
            "Libro:",
            value=campos['libro']
        )
        campos['certificado_existencia'] = st.text_input(
            "Año del certificado:",
            value=campos['certificado_existencia']
        )
    
    st.markdown("### 👨‍⚖️ Información del Representante Legal")
    st.markdown("---")
    col1, col2 = st.columns(2)
    
    with col1:
        campos['nombre_representante'] = st.text_input(
            "Nombre del representante legal:",
            value=campos['nombre_representante']
        )
        campos['cedula_representante'] = st.text_input(
            "Cédula del representante:",
            value=campos['cedula_representante']
        )
        campos['ciudad_representante'] = st.text_input(
            "Ciudad del representante:",
            value=campos['ciudad_representante']
        )
    
    with col2:
        campos['tarjeta_profesional'] = st.text_input(
            "Tarjeta profesional:",
            value=campos['tarjeta_profesional']
        )
        campos['consejo_judicatura'] = st.text_input(
            "Consejo de judicatura:",
            value=campos['consejo_judicatura']
        )
        campos['ciudad_representante_firma'] = st.text_input(
            "Ciudad para firma del representante:",
            value=campos['ciudad_representante_firma']
        )
    
    st.markdown("### ⚖️ Información del Caso")
    st.markdown("---")
    col1, col2 = st.columns(2)
    
    with col1:
        campos['entidad_demandada'] = st.text_input(
            "Entidad demandada:",
            value=campos['entidad_demandada']
        )
        campos['oficio_1'] = st.text_input(
            "Número del oficio 1:",
            value=campos['oficio_1']
        )
        campos['fecha_oficio_1'] = st.text_input(
            "Fecha del oficio 1:",
            value=campos['fecha_oficio_1']
        )
        campos['jefe_regional'] = st.text_input(
            "Jefe regional:",
            value=campos['jefe_regional']
        )
    
    with col2:
        campos['oficio_2'] = st.text_input(
            "Número del oficio 2:",
            value=campos['oficio_2']
        )
        campos['fecha_oficio_2'] = st.text_input(
            "Fecha del oficio 2:",
            value=campos['fecha_oficio_2']
        )
        campos['directora_sanidad'] = st.text_input(
            "Directora de sanidad:",
            value=campos['directora_sanidad']
        )
    
    st.markdown("### 💼 Información Laboral")
    st.markdown("---")
    col1, col2, col3 = st.columns(3)
    
    with col1:
        campos['cargo_laboral'] = st.text_input(
            "Cargo laboral:",
            value=campos['cargo_laboral']
        )
    
    with col2:
        campos['fecha_inicio_laboral'] = st.text_input(
            "Fecha de inicio:",
            value=campos['fecha_inicio_laboral']
        )
    
    with col3:
        campos['fecha_fin_laboral'] = st.text_input(
            "Fecha de fin:",
            value=campos['fecha_fin_laboral']
        )
    
    st.markdown("---")
    
    # Botón para generar documento
    st.markdown("---")
    if st.button("📄 Generar Documento de Poder", type="primary", use_container_width=True):
        try:
            doc = generar_documento_poder(campos)
            
            # Guardar en buffer para descarga
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            st.success("✅ Documento de poder generado exitosamente!")
            
            # Botón de descarga
            st.download_button(
                label="📥 Descargar Poder en Word",
                data=buffer,
                file_name="poder_especial.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
        except Exception as e:
            st.error(f"❌ Error al generar el documento: {str(e)}")
    
    # Vista previa del documento
    if st.checkbox("👁️ Mostrar vista previa del documento"):
        st.subheader("📋 Vista Previa del Poder")
        doc_preview = generar_documento_poder(campos)
        texto_preview = ""
        for paragraph in doc_preview.paragraphs:
            texto_preview += paragraph.text + "\n\n"
        
        st.text_area(
            "Contenido del documento:",
            value=texto_preview,
            height=400,
            disabled=True
        ) 