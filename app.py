# app.py - Asistente Jurídico Contrato Realidad AI

import streamlit as st
import io
import os
from utils.exportar import generar_docx_concepto
from utils.resumen import generar_resumen
from utils.viabilidad import evaluar_viabilidad
from utils.por_secciones import generar_seccion
from utils.rag import generar_resumen_con_rag, evaluar_viabilidad_con_rag, generar_seccion_con_rag
from utils.vector_rag import generar_resumen_vector_rag, evaluar_viabilidad_vector_rag, generar_seccion_vector_rag
from utils.knowledge_manager import render_knowledge_manager
from utils.poder import render_poder_module
from utils.transcripcion import render_transcripcion_module, render_transcripcion_inline
from utils.expediente import render_cargar_expediente
from utils.documento_referencia import generar_seccion_con_referencia
from docx import Document

# Ruta del logo (con manejo de error si no existe)
logo_path = os.path.join('assets', 'logo_conde_abogados.png')
if not os.path.exists(logo_path):
    logo_path = None  # No mostrar logo si no existe

# Configuración inicial de la app
st.set_page_config(
    page_title="Asistente Jurídico - Contrato Realidad",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS personalizado para mejorar el diseño
st.markdown("""
<style>
    /* Estilos generales */
    .main {
        padding-top: 2rem;
    }
    
    /* Cards modernos */
    .card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 1.5rem;
        border-radius: 15px;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
        margin: 1rem 0;
        color: white;
    }
    
    .card-info {
        background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%);
        padding: 1.5rem;
        border-radius: 15px;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
        margin: 1rem 0;
        color: white;
    }
    
    .card-success {
        background: linear-gradient(135deg, #4facfe 0%, #00f2fe 100%);
        padding: 1.5rem;
        border-radius: 15px;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
        margin: 1rem 0;
        color: white;
    }
    
    /* Indicadores de progreso */
    .progress-container {
        display: flex;
        justify-content: space-between;
        margin: 2rem 0;
        padding: 1rem;
        background: #f8f9fa;
        border-radius: 10px;
    }
    
    .progress-step {
        flex: 1;
        text-align: center;
        padding: 0.5rem;
        position: relative;
    }
    
    .progress-step.active {
        color: #667eea;
        font-weight: bold;
    }
    
    .progress-step.completed {
        color: #28a745;
    }
    
    /* Botones mejorados */
    .stButton > button {
        border-radius: 10px;
        transition: all 0.3s ease;
        font-weight: 600;
    }
    
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 8px rgba(0, 0, 0, 0.2);
    }
    
    /* Sidebar mejorado */
    .css-1d391kg {
        background: linear-gradient(180deg, #667eea 0%, #764ba2 100%);
    }
    
    /* Texto destacado */
    .highlight {
        background: linear-gradient(120deg, #84fab0 0%, #8fd3f4 100%);
        padding: 0.2rem 0.5rem;
        border-radius: 5px;
        font-weight: 600;
    }
    
    /* Separadores mejorados */
    hr {
        border: none;
        height: 2px;
        background: linear-gradient(90deg, transparent, #667eea, transparent);
        margin: 2rem 0;
    }
    
    /* Mejoras en text areas */
    .stTextArea > div > div > textarea {
        border-radius: 10px;
        border: 2px solid #e0e0e0;
    }
    
    .stTextArea > div > div > textarea:focus {
        border-color: #667eea;
        box-shadow: 0 0 0 3px rgba(102, 126, 234, 0.1);
    }
    
    /* Mejoras en selectbox */
    .stSelectbox > div > div {
        border-radius: 10px;
    }
</style>
""", unsafe_allow_html=True)

def render_progress_indicator(current_phase: int, total_phases: int = 5):
    """Renderiza un indicador de progreso visual"""
    phases = [
        ("📥", "Carga de Datos", 0),
        ("📝", "Hechos y Resumen", 1),
        ("⚖️", "Evaluación", 2),
        ("📋", "Poder", 3),
        ("📜", "Demanda", 4)
    ]
    
    cols = st.columns(5)
    for idx, (icon, name, phase_num) in enumerate(phases):
        with cols[idx]:
            if phase_num < current_phase:
                status = "✅"
                color = "#28a745"
            elif phase_num == current_phase:
                status = "🔄"
                color = "#667eea"
            else:
                status = "⏳"
                color = "#6c757d"
            
            st.markdown(f"""
            <div style="text-align: center; padding: 1rem; background: {'#e3f2fd' if phase_num == current_phase else '#f5f5f5'}; 
                        border-radius: 10px; border: {'3px solid #667eea' if phase_num == current_phase else '1px solid #ddd'};">
                <div style="font-size: 2rem;">{status}</div>
                <div style="font-weight: {'bold' if phase_num == current_phase else 'normal'}; color: {color}; margin-top: 0.5rem; font-size: 0.9rem;">
                    {name}
                </div>
            </div>
            """, unsafe_allow_html=True)

def render_phase_card(title: str, icon: str, content: str):
    """Renderiza una tarjeta de fase"""
    st.markdown(f"""
    <div class="card">
        <h2 style="color: white; margin-bottom: 1rem;">{icon} {title}</h2>
        <p style="color: white; font-size: 1.1rem;">{content}</p>
    </div>
    """, unsafe_allow_html=True)

# Sidebar para navegación y configuración
with st.sidebar:
    if logo_path and os.path.exists(logo_path):
        try:
            st.image(logo_path, width=220)
        except Exception:
            st.markdown("### ⚖️ Asistente Jurídico")
    else:
        st.markdown("### ⚖️ Asistente Jurídico")
    
    st.markdown("---")
    st.header("🧭 Navegación")
    
    # Selector de página
    page = st.selectbox(
        "Seleccionar módulo:",
        ["🏠 Asistente Jurídico", "📋 Módulo de Poder", "🎤 Transcripción de Audios", "📚 Gestor de Conocimiento"],
        help="Elige la funcionalidad que deseas usar"
    )
    
    st.markdown("---")
    
    # Configuración RAG (solo para la página del asistente)
    if page == "🏠 Asistente Jurídico":
        st.header("⚙️ Configuración RAG")
        rag_mode = st.selectbox(
            "Modo de RAG:",
            ["Sin RAG", "RAG Básico", "RAG Vectorial"],
            help="RAG Básico: Búsqueda por palabras clave\nRAG Vectorial: Búsqueda semántica con embeddings"
        )
        
        if rag_mode != "Sin RAG":
            st.success(f"✅ Usando {rag_mode}")
            st.markdown("**Beneficios:**")
            st.markdown("• Respuestas fundamentadas")
            st.markdown("• Referencias legales")
            st.markdown("• Información actualizada")
        else:
            st.info("💡 Activa RAG para respuestas más precisas")
        
        st.markdown("---")
        
        # Información del caso
        if "hechos" in st.session_state and st.session_state.hechos:
            st.subheader("📄 Información del Caso")
            st.caption("Hechos ingresados")
            if st.button("🔄 Reiniciar Caso"):
                for key in list(st.session_state.keys()):
                    if key not in ["rag_mode_sidebar"]:
                        del st.session_state[key]
                st.rerun()

# Página principal del asistente jurídico
if page == "🏠 Asistente Jurídico":
    # Título principal con estilo
    st.markdown("""
    <div style="text-align: center; padding: 2rem 0; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); 
                border-radius: 15px; margin-bottom: 2rem;">
        <h1 style="color: white; margin: 0;">⚖️ Asistente Jurídico</h1>
        <p style="color: white; font-size: 1.2rem; margin-top: 0.5rem;">Contrato Realidad - Análisis y Generación de Documentos Legales</p>
    </div>
    """, unsafe_allow_html=True)

    # Inicialización del estado de sesión
    if "fase" not in st.session_state:
        st.session_state.fase = 0
    if "metodo_carga" not in st.session_state:
        st.session_state.metodo_carga = None
    if "hechos" not in st.session_state:
        st.session_state.hechos = ""
    if "resumen" not in st.session_state:
        st.session_state.resumen = ""
    if "resumen_generado" not in st.session_state:
        st.session_state.resumen_generado = False
    if "concepto" not in st.session_state:
        st.session_state.concepto = ""
    if "demanda" not in st.session_state:
        st.session_state.demanda = ""
    if "demanda_generada" not in st.session_state:
        st.session_state.demanda_generada = False
    if "nombre_abogado" not in st.session_state:
        st.session_state.nombre_abogado = ""
    if "seccion_actual" not in st.session_state:
        st.session_state.seccion_actual = 0
    if "secciones_demanda" not in st.session_state:
        st.session_state.secciones_demanda = {
            "I. Hechos": "",
            "II. Peticiones": "",
            "III. Petición Final": "",
            "IV. Fundamentos de derecho": "",
            "V. Normatividad y jurisprudencia aplicable al caso": "",
            "VI. Relación de medios probatorios": "",
            "VII. Cuantía": "",
            "VIII. Propuesta de fórmula de conciliación": "",
            "IX. Competencia": "",
            "X. Manifestación": "",
            "XI. Anexos": "",
            "XII. Notificaciones": ""
        }
    # Cargar patrones de referencia automáticamente si existe el archivo
    if "patrones_referencia" not in st.session_state:
        import json
        import os
        patrones_file = "patrones_referencia.json"
        if os.path.exists(patrones_file):
            try:
                with open(patrones_file, 'r', encoding='utf-8') as f:
                    st.session_state.patrones_referencia = json.load(f)
            except Exception:
                st.session_state.patrones_referencia = {}
        else:
            st.session_state.patrones_referencia = {}
    if "poder_generado" not in st.session_state:
        st.session_state.poder_generado = False

    # Obtener modo RAG del sidebar
    rag_mode = st.sidebar.selectbox(
        "Modo de RAG:",
        ["Sin RAG", "RAG Básico", "RAG Vectorial"],
        key="rag_mode_sidebar"
    )

    # Indicador de progreso (solo mostrar si no estamos en fase intermedia)
    if st.session_state.fase != 0.5:
        render_progress_indicator(int(st.session_state.fase))

    # -------------------------------------------
    # FASE 0: SELECCIÓN DE MÉTODO DE CARGA
    # -------------------------------------------
    if st.session_state.fase == 0:
        render_phase_card(
            "Fase 0: Selección de Método de Carga",
            "📥",
            "Elige cómo deseas cargar la información del caso: transcripción de entrevista o expediente"
        )
        
        st.markdown("### 🎯 ¿Cómo deseas iniciar el caso?")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("""
            <div style="padding: 2rem; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); 
                        border-radius: 15px; text-align: center; color: white; margin-bottom: 1rem;">
                <h2 style="color: white; margin-bottom: 1rem;">🎤 Transcripción de Entrevista</h2>
                <p style="color: white; font-size: 1.1rem;">
                    Sube un archivo de audio de la entrevista con el cliente y conviértelo en texto automáticamente
                </p>
            </div>
            """, unsafe_allow_html=True)
            
            if st.button("🎤 Usar Transcripción", type="primary", use_container_width=True):
                st.session_state.metodo_carga = "transcripcion"
                st.session_state.fase = 0.5  # Fase intermedia para transcripción
                st.rerun()
        
        with col2:
            st.markdown("""
            <div style="padding: 2rem; background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%); 
                        border-radius: 15px; text-align: center; color: white; margin-bottom: 1rem;">
                <h2 style="color: white; margin-bottom: 1rem;">📄 Cargar Expediente</h2>
                <p style="color: white; font-size: 1.1rem;">
                    Sube un archivo PDF o pega el texto del expediente para extraer la información del caso
                </p>
            </div>
            """, unsafe_allow_html=True)
            
            if st.button("📄 Usar Expediente", type="primary", use_container_width=True):
                st.session_state.metodo_carga = "expediente"
                st.session_state.fase = 0.5  # Fase intermedia para expediente
                st.rerun()
        
        st.markdown("---")
        st.info("💡 También puedes escribir los hechos manualmente en la siguiente fase si lo prefieres.")

    # -------------------------------------------
    # FASE 0.5: PROCESAMIENTO DE DATOS
    # -------------------------------------------
    elif st.session_state.fase == 0.5:
        if st.session_state.metodo_carga == "transcripcion":
            render_phase_card(
                "Transcripción de Entrevista",
                "🎤",
                "Sube el archivo de audio de la entrevista para transcribirlo automáticamente"
            )
            
            texto_transcrito = render_transcripcion_inline()
            
            if texto_transcrito:
                st.session_state.hechos = texto_transcrito
                st.success("✅ Transcripción completada. Los hechos han sido cargados.")
                
                col1, col2 = st.columns(2)
                with col1:
                    if st.button("⬅️ Volver", use_container_width=True):
                        st.session_state.fase = 0
                        st.session_state.metodo_carga = None
                        st.rerun()
                
                with col2:
                    if st.button("➡️ Continuar con estos hechos", type="primary", use_container_width=True):
                        st.session_state.fase = 1
                        st.rerun()
        
        elif st.session_state.metodo_carga == "expediente":
            render_phase_card(
                "Carga de Expediente",
                "📄",
                "Carga el expediente del caso (PDF o texto) para extraer la información"
            )
            
            texto_expediente = render_cargar_expediente()
            
            if texto_expediente:
                st.session_state.hechos = texto_expediente
                st.success("✅ Expediente cargado. Los hechos han sido extraídos.")
                
                col1, col2 = st.columns(2)
                with col1:
                    if st.button("⬅️ Volver", use_container_width=True):
                        st.session_state.fase = 0
                        st.session_state.metodo_carga = None
                        st.rerun()
                
                with col2:
                    if st.button("➡️ Continuar con este expediente", type="primary", use_container_width=True):
                        st.session_state.fase = 1
                        st.rerun()

    # -------------------------------------------
    # FASE 1: INGRESO DE HECHOS Y RESUMEN
    # -------------------------------------------
    elif st.session_state.fase == 1:
        render_phase_card(
            "Fase 1: Ingreso de Hechos y Resumen Jurídico",
            "📝",
            "Revisa y edita los hechos del caso (si fueron cargados) o ingrésalos manualmente, luego genera un resumen técnico jurídico"
        )

        # Mostrar método de carga usado
        if st.session_state.metodo_carga:
            metodo_nombre = "Transcripción de entrevista" if st.session_state.metodo_carga == "transcripcion" else "Expediente cargado"
            st.info(f"ℹ️ Hechos cargados desde: **{metodo_nombre}**. Puedes editarlos a continuación.")

        st.markdown("### ✍️ Hechos del caso")
        hechos = st.text_area(
            "Hechos del caso:",
            height=300,
            value=st.session_state.hechos,
            placeholder="Ejemplo: El trabajador prestó servicios desde enero de 2011 hasta septiembre de 2019, recibía órdenes directas de su supervisor, trabajaba de lunes a viernes en horario fijo, y recibía pagos mensuales, pero el contrato estaba bajo la figura de prestación de servicios..."
        )
        
        # Opción para cambiar el método de carga
        if st.button("🔄 Cambiar método de carga"):
            st.session_state.fase = 0
            st.session_state.metodo_carga = None
            st.rerun()

        col1, col2 = st.columns([1, 1])
        with col1:
            if st.button("🔍 Generar Resumen Técnico", type="primary", use_container_width=True):
                if hechos.strip() == "":
                    st.warning("⚠️ Por favor, ingresa los hechos antes de continuar.")
                else:
                    with st.spinner("🤖 Generando resumen técnico con IA..."):
                        # Usar RAG según el modo seleccionado
                        if rag_mode == "RAG Básico":
                            resumen = generar_resumen_con_rag(hechos)
                        elif rag_mode == "RAG Vectorial":
                            resumen = generar_resumen_vector_rag(hechos)
                        else:
                            resumen = generar_resumen(hechos)
                            
                    st.session_state.resumen = resumen
                    st.session_state.hechos = hechos
                    st.session_state.resumen_generado = True
                    st.success("✅ Resumen generado con éxito!")
                    
                    with st.expander("📌 Ver Resumen Técnico del Caso", expanded=True):
                        st.markdown(resumen)

        with col2:
            if st.button("➡️ Continuar a Fase 2", type="primary", use_container_width=True):
                if hechos.strip() == "":
                    st.warning("⚠️ Debes ingresar los hechos antes de continuar.")
                elif not st.session_state.resumen_generado:
                    st.warning("⚠️ Primero debes generar el resumen técnico.")
                else:
                    st.session_state.fase = 2
                    st.rerun()

    # -------------------------------------------
    # FASE 2: CONCEPTO DE VIABILIDAD
    # -------------------------------------------
    elif st.session_state.fase == 2:
        render_phase_card(
            "Fase 2: Evaluación Jurídica del Caso",
            "⚖️",
            "Evalúa la viabilidad jurídica del caso y genera un concepto técnico fundamentado"
        )

        st.markdown("### 📄 Hechos del caso")
        st.info(st.session_state.hechos)

        if st.button("📘 Emitir Concepto Jurídico", type="primary", use_container_width=True):
            with st.spinner("🤖 Evaluando viabilidad jurídica con IA..."):
                # Usar RAG según el modo seleccionado
                if rag_mode == "RAG Básico":
                    concepto = evaluar_viabilidad_con_rag(st.session_state.hechos)
                elif rag_mode == "RAG Vectorial":
                    concepto = evaluar_viabilidad_vector_rag(st.session_state.hechos)
                else:
                    concepto = evaluar_viabilidad(st.session_state.hechos)
                    
            st.session_state.concepto = concepto
            st.success("✅ Concepto generado exitosamente!")
            
            with st.expander("📑 Ver Concepto Jurídico", expanded=True):
                st.markdown(concepto)

            # Descargar concepto
            doc = generar_docx_concepto(st.session_state.hechos, concepto)
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                label="📥 Descargar Concepto en Word",
                data=buffer,
                file_name="concepto_viabilidad.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

        col1, col2 = st.columns([1, 1])
        with col1:
            if st.button("⬅️ Volver a Fase 1", use_container_width=True):
                st.session_state.fase = 1
                st.rerun()
        
        with col2:
            if st.button("➡️ Continuar a Fase 3", type="primary", use_container_width=True):
                if not st.session_state.concepto:
                    st.warning("⚠️ Primero debes generar el concepto jurídico.")
                else:
                    st.session_state.fase = 3
                    st.rerun()

    # -------------------------------------------
    # FASE 3: GENERACIÓN DE PODER
    # -------------------------------------------
    elif st.session_state.fase == 3:
        render_phase_card(
            "Fase 3: Generación de Poder Especial",
            "📋",
            "Genera el poder especial para representar al cliente en el proceso judicial"
        )
        
        st.info("""
        **¿Qué es un Poder Especial?**
        
        Un poder especial es un documento legal que autoriza a un abogado o firma jurídica 
        para representar a un cliente en un proceso judicial específico. En este caso, 
        se genera un poder para presentar un medio de control de nulidad y restablecimiento del derecho.
        """)
        
        # Mostrar información del caso
        st.markdown("### 📄 Información del Caso")
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown(f"""
            <div style="padding: 1rem; background: #f8f9fa; border-radius: 10px; margin-bottom: 1rem;">
                <strong>Poderdante:</strong> {st.session_state.nombre_abogado if st.session_state.nombre_abogado else "Por definir"}<br>
                <strong>Caso:</strong> Contrato Realidad<br>
                <strong>Entidad:</strong> Ministerio de Defensa - Policía Nacional
            </div>
            """, unsafe_allow_html=True)
        
        with col2:
            st.markdown(f"""
            <div style="padding: 1rem; background: #f8f9fa; border-radius: 10px; margin-bottom: 1rem;">
                <strong>Tipo de Proceso:</strong> Nulidad y Restablecimiento del Derecho<br>
                <strong>Objetivo:</strong> Reconocimiento de vínculo laboral<br>
                <strong>Período:</strong> 2011 - 2019
            </div>
            """, unsafe_allow_html=True)
        
        st.markdown("---")
        
        col1, col2, col3 = st.columns([1, 1, 1])
        with col1:
            if st.button("⬅️ Volver a Fase 2", use_container_width=True):
                st.session_state.fase = 2
                st.rerun()
        
        with col2:
            if st.button("📋 Ir al Módulo de Poder", type="primary", use_container_width=True):
                st.session_state.poder_generado = True
                st.session_state.fase = 4
                st.rerun()
        
        with col3:
            if st.button("⏭️ Saltar y Continuar", use_container_width=True):
                st.session_state.fase = 4
                st.rerun()

    # -------------------------------------------
    # FASE 4: REDACCIÓN DE LA DEMANDA
    # -------------------------------------------
    elif st.session_state.fase == 4:
        render_phase_card(
            "Fase 4: Redacción Guiada de la Demanda",
            "📜",
            "Redacta la demanda laboral sección por sección con asistencia de IA"
        )

        if not (st.session_state.hechos and st.session_state.resumen_generado and st.session_state.concepto):
            st.error("⚠️ Debes completar las fases anteriores antes de redactar la demanda.")
            if st.button("⬅️ Volver al Inicio"):
                st.session_state.fase = 1
                st.rerun()
        else:
            # Mostrar información del poder si fue generado
            if st.session_state.poder_generado:
                st.success("✅ Poder especial generado anteriormente")
            
            
            # Solicitar el nombre del abogado
            nombre = st.text_input(
                "👨‍⚖️ Ingresa el nombre del abogado supervisor:",
                value=st.session_state.nombre_abogado,
                placeholder="Ej: Dr. Juan Pérez"
            )

            if not nombre.strip():
                st.warning("⚠️ Por favor, ingresa el nombre del abogado para continuar.")
            else:
                st.session_state.nombre_abogado = nombre.strip()

                # Redacción guiada sección por sección
                secciones = list(st.session_state.secciones_demanda.keys())
                seccion = secciones[st.session_state.seccion_actual]
                
                st.markdown(f"""
                <div style="padding: 1.5rem; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); 
                            border-radius: 15px; margin: 1rem 0;">
                    <h3 style="color: white; margin: 0;">✒️ Sección {st.session_state.seccion_actual + 1} de {len(secciones)}: {seccion}</h3>
                </div>
                """, unsafe_allow_html=True)

                if st.session_state.secciones_demanda[seccion] == "":
                    with st.spinner(f"🤖 Redactando la sección: {seccion}..."):
                        # Si hay patrones de referencia cargados automáticamente, usarlos
                        if st.session_state.patrones_referencia:
                            texto_generado = generar_seccion_con_referencia(
                                seccion,
                                st.session_state.hechos,
                                st.session_state.resumen,
                                st.session_state.concepto,
                                st.session_state.patrones_referencia
                            )
                        # Usar RAG según el modo seleccionado
                        elif rag_mode == "RAG Básico":
                            texto_generado = generar_seccion_con_rag(
                                seccion,
                                st.session_state.hechos,
                                st.session_state.resumen,
                                st.session_state.concepto
                            )
                        elif rag_mode == "RAG Vectorial":
                            texto_generado = generar_seccion_vector_rag(
                                seccion,
                                st.session_state.hechos,
                                st.session_state.resumen,
                                st.session_state.concepto
                            )
                        else:
                            texto_generado = generar_seccion(
                                seccion,
                                st.session_state.hechos,
                                st.session_state.resumen,
                                st.session_state.concepto
                            )
                            
                    st.session_state.secciones_demanda[seccion] = texto_generado

                st.text_area(
                    "📄 Redacción sugerida por el sistema:",
                    value=st.session_state.secciones_demanda[seccion],
                    height=300,
                    key=f"text_area_{seccion}"
                )

                acuerdo = st.radio(
                    "¿Estás de acuerdo con esta sección?",
                    ["Sí", "No"],
                    key=f"acuerdo_{seccion}",
                    horizontal=True
                )

                if acuerdo == "Sí":
                    col1, col2 = st.columns([1, 1])
                    with col1:
                        if st.session_state.seccion_actual > 0:
                            if st.button("⬅️ Sección Anterior", use_container_width=True):
                                st.session_state.seccion_actual -= 1
                                st.rerun()
                    
                    with col2:
                        if st.session_state.seccion_actual < len(secciones) - 1:
                            if st.button("➡️ Siguiente Sección", type="primary", use_container_width=True):
                                st.session_state.seccion_actual += 1
                                st.rerun()
                        else:
                            st.success("🎉 ¡Demanda redactada completamente!")
                            
                            # Descargar la demanda cuando todas las secciones han sido aceptadas
                            if st.button("📥 Descargar Demanda Completa en Word", type="primary", use_container_width=True):
                                buffer = io.BytesIO()
                                doc = Document()
                                doc.add_heading("Demanda por Contrato Realidad", level=1)
                                doc.add_paragraph(f"Abogado supervisor: {st.session_state.nombre_abogado}")
                                doc.add_paragraph("")
                                for titulo, contenido in st.session_state.secciones_demanda.items():
                                    doc.add_heading(titulo, level=2)
                                    doc.add_paragraph(contenido)
                                    doc.add_paragraph("")
                                doc.save(buffer)
                                buffer.seek(0)

                                st.download_button(
                                    label="📄 Descargar Demanda en Word",
                                    data=buffer,
                                    file_name="demanda_contrato_realidad.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    use_container_width=True
                                )
                else:
                    comentario = st.text_area(
                        "✍️ ¿Qué ajustes deseas hacer?",
                        key=f"comentario_{seccion}",
                        placeholder="Ej: Hacer más énfasis en la subordinación, mencionar jurisprudencia específica..."
                    )
                    if st.button("🔁 Reescribir esta sección con los comentarios", type="primary", use_container_width=True):
                        with st.spinner("🤖 Reescribiendo sección..."):
                            # Si hay patrones de referencia cargados automáticamente, usarlos
                            if st.session_state.patrones_referencia:
                                nueva_redaccion = generar_seccion_con_referencia(
                                    seccion,
                                    st.session_state.hechos,
                                    st.session_state.resumen,
                                    st.session_state.concepto,
                                    st.session_state.patrones_referencia,
                                    comentario_usuario=comentario
                                )
                            # Usar RAG según el modo seleccionado
                            elif rag_mode == "RAG Básico":
                                nueva_redaccion = generar_seccion_con_rag(
                                    seccion,
                                    st.session_state.hechos,
                                    st.session_state.resumen,
                                    st.session_state.concepto,
                                    comentario_usuario=comentario
                                )
                            elif rag_mode == "RAG Vectorial":
                                nueva_redaccion = generar_seccion_vector_rag(
                                    seccion,
                                    st.session_state.hechos,
                                    st.session_state.resumen,
                                    st.session_state.concepto,
                                    comentario_usuario=comentario
                                )
                            else:
                                nueva_redaccion = generar_seccion(
                                    seccion,
                                    st.session_state.hechos,
                                    st.session_state.resumen,
                                    st.session_state.concepto,
                                    comentario_usuario=comentario
                                )
                                
                        st.session_state.secciones_demanda[seccion] = nueva_redaccion
                        st.rerun()

# Página del módulo de poder
elif page == "📋 Módulo de Poder":
    render_poder_module()

# Página del módulo de transcripción
elif page == "🎤 Transcripción de Audios":
    render_transcripcion_module()

# Página del gestor de conocimiento
elif page == "📚 Gestor de Conocimiento":
    st.title("📚 Gestor de Conocimiento Legal")
    render_knowledge_manager()
